import docx
from docx import Document
from docx.shared import RGBColor
import string
import re


filename = "SRS_ACE_Pump_X00.docx"


def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for paras in doc.paragraphs:
        for run in paras.runs:
            if run.font.color.rgb == RGBColor(255, 000, 000):
                fullText.append(run.text)
    return fullText





fullText = readtxt("SRS_ACE_Pump_X00.docx")

s = ''.join(fullText)
print(s)

#o = s.split(']\w+')
#print(o)


#print str(fullText)[1:-1]
#print(fullText)
#print(fullText[5])

report = Document()
#paragraph = report.add_paragraph(fullText[1])
#paragraph = report.add_paragraph(fullText)
#paragraph = report.add_paragraph(s)
#report.save('report1.docx')


w = (s.replace (']', ']\n'))
paragraph = report.add_paragraph()
runner = paragraph.add_run("\nSRS Ace Pump Document")
runner.bold = True #makes the header bold
paragraph = report.add_paragraph(w)

runner = paragraph.add_run("\nPRS new Pump Document")
runner.bold = True

filename2 = "`PRS_new_pump.docx"
fullText2 = readtxt("PRS_new_pump.docx")

b = ''.join(fullText2)
c = (b.replace (']', ']\n'))
print(w)

paragraph = report.add_paragraph(c)


runner = paragraph.add_run("\nHTP_new_pump.docx")
runner.bold = True

filename3 = "HTP_new_pump.docx"
fullText3 = readtxt("HTP_new_pump.docx")

d = ''.join(fullText3)
e = (d.replace (']', ']\n'))
print(e)

paragraph = report.add_paragraph(e)


'''report.add_heading('System Details', 0)
# Table data in a form of list
data = (
    (s, w, e),
    (s, w, e),
    (s, w, e)
)
# Creating a table object
table = report.add_table(rows=1, cols=3)
# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Entry Tag'
row[1].text = 'Info'
row[2].text = 'Ending Tag'
# Adding data from the list to the table
for id, info, end in data:
    # Adding a row and then adding data in it.
    row = table.add_row().cells
    # Converting id to string as table can only take string input
    row[0].text = id
    row[1].text = info
    row[2].text = end

table.style = 'Colorful List'
'''

report.save('report1.docx')
